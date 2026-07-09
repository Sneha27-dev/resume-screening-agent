"""
parser.py

Reads resume files (PDF, DOCX, TXT) and converts them into clean text.

Features
--------
- Supports PDF, DOCX and TXT
- Cleans extracted text
- Skips unsupported files
- Handles parsing errors gracefully
- Detailed logging
"""

from pathlib import Path

from utils import clean_text, logger


# ---------------------------------------------------------------------
# TXT Reader
# ---------------------------------------------------------------------

def read_txt(path: Path) -> str:
    """Read plain text file."""
    return path.read_text(
        encoding="utf-8",
        errors="ignore",
    )


# ---------------------------------------------------------------------
# PDF Reader
# ---------------------------------------------------------------------

def read_pdf(path: Path) -> str:
    """Extract text from PDF."""

    try:
        import pdfplumber

    except ImportError as e:

        raise ImportError(
            "Missing dependency: pdfplumber\n"
            "Install using:\n"
            "pip install pdfplumber"
        ) from e

    pages = []

    with pdfplumber.open(path) as pdf:

        for page in pdf.pages:

            pages.append(page.extract_text() or "")

    return "\n".join(pages)


# ---------------------------------------------------------------------
# DOCX Reader
# ---------------------------------------------------------------------

def read_docx(path: Path) -> str:
    """Extract text from DOCX."""

    try:
        import docx

    except ImportError as e:

        raise ImportError(
            "Missing dependency: python-docx\n"
            "Install using:\n"
            "pip install python-docx"
        ) from e

    document = docx.Document(path)

    paragraphs = []

    # Normal paragraphs
    for paragraph in document.paragraphs:
        paragraphs.append(paragraph.text)

    # Tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return "\n".join(paragraphs)


# ---------------------------------------------------------------------
# Supported Readers
# ---------------------------------------------------------------------

READERS = {
    ".txt": read_txt,
    ".pdf": read_pdf,
    ".docx": read_docx,
}


# ---------------------------------------------------------------------
# Parse One Resume
# ---------------------------------------------------------------------

def parse_resume(path: Path) -> str:
    """
    Parse one resume into cleaned plain text.
    """

    if not path.exists():

        logger.error(f"File not found: {path}")

        return ""

    extension = path.suffix.lower()

    reader = READERS.get(extension)

    if reader is None:

        logger.warning(
            f"Unsupported file type: {path.name}"
        )

        return ""

    logger.info(f"Reading: {path.name}")

    try:

        raw_text = reader(path)

        cleaned = clean_text(raw_text)

        if cleaned.strip():

            logger.info(
                f"Successfully parsed: {path.name}"
            )

            return cleaned

        logger.warning(
            f"No extractable text found in {path.name}"
        )

        return ""

    except Exception as e:

        logger.warning(
            f"Failed to parse {path.name}: {e}"
        )

        return ""


# ---------------------------------------------------------------------
# Load Multiple Resumes
# ---------------------------------------------------------------------

def load_resumes(folder: Path) -> dict:
    """
    Parse every supported resume inside a folder.

    Returns
    -------
    dict
        {
            filename: extracted_text
        }
    """

    from config import SUPPORTED_EXTENSIONS

    resumes = {}

    if not folder.exists():

        logger.error(
            f"Resume folder not found: {folder}"
        )

        return resumes

    logger.info(f"Scanning folder: {folder}")

    files = sorted(folder.iterdir())

    if not files:

        logger.warning("No resume files found.")

        return resumes

    for path in files:

        if path.name.startswith("."):
            continue

        if not path.is_file():
            continue

        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:

            logger.info(
                f"Skipping unsupported file: {path.name}"
            )

            continue

        text = parse_resume(path)

        if text:

            resumes[path.name] = text

    logger.info(
        f"Successfully loaded {len(resumes)} resume(s)."
    )

    return resumes